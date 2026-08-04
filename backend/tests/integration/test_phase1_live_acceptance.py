from __future__ import annotations

import json
import os
from dataclasses import dataclass, replace
from datetime import UTC, datetime, timedelta
from io import BytesIO
from urllib.parse import urlparse
from uuid import UUID, uuid4

import boto3  # type: ignore[import-untyped]
import fitz
import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from openpyxl import Workbook
from redis import Redis
from sqlalchemy import delete, select
from sqlalchemy.engine import make_url

from knowagent.api.app import create_app
from knowagent.common.lifecycle import PublicationStatus
from knowagent.documents.application.chunking import ChunkingConfig, StructureAwareChunker
from knowagent.documents.application.processor import (
    ChunkManifest,
    IngestionProcessor,
    IngestionRecoveryService,
)
from knowagent.documents.domain.ingestion import DocumentVersionStatus, IngestionStatus
from knowagent.documents.domain.models import SourceType
from knowagent.documents.infrastructure.parsers import ParserLimits
from knowagent.documents.infrastructure.parsers.registry import ParserRegistry
from knowagent.documents.infrastructure.sqlalchemy_models import (
    DocumentRecord,
    DocumentVersionRecord,
)
from knowagent.documents.infrastructure.sqlalchemy_repository import (
    SqlAlchemyIngestionRepository,
)
from knowagent.identity.domain.models import AccountRole, AccountSource, AccountStatus
from knowagent.identity.infrastructure.passwords import Argon2PasswordHasher
from knowagent.identity.infrastructure.sqlalchemy_models import AccountRecord, AuditLogRecord
from knowagent.knowledge.application.publication import KnowledgePublicationService
from knowagent.knowledge.domain.models import KnowledgeChunkDraft
from knowagent.knowledge.infrastructure.sqlalchemy_repository import (
    SqlAlchemyKnowledgeRepository,
)
from knowagent.platform.object_store import ObjectStoreError, S3ObjectStore
from knowagent.platform.settings import ObjectStorageSettings, Settings
from knowagent.systems.infrastructure.sqlalchemy_models import (
    AccountSystemRoleRecord,
    BusinessSystemRecord,
)
from knowagent.worker.celery_app import build_celery_app
from knowagent.worker.dispatcher import CeleryIngestionDispatcher

pytestmark = [
    pytest.mark.integration,
    pytest.mark.skipif(
        os.getenv("KNOWAGENT_RUN_PHASE1_ACCEPTANCE") != "1",
        reason="set KNOWAGENT_RUN_PHASE1_ACCEPTANCE=1 to run live infrastructure acceptance",
    ),
]

PASSWORD = "Temporary1!"


@dataclass(frozen=True, slots=True)
class Sample:
    filename: str
    media_type: str
    content: bytes
    source_type: SourceType


def _make_docx() -> bytes:
    target = BytesIO()
    document = DocxDocument()
    document.add_heading("ESB Integration Guide", level=1)
    document.add_paragraph("Request client credentials before connecting.")
    document.add_heading("Timeout", level=2)
    document.add_paragraph("The default timeout is 30 seconds.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "parameter"
    table.cell(0, 1).text = "meaning"
    table.cell(1, 0).text = "timeout"
    table.cell(1, 1).text = "request timeout"
    document.save(target)
    return target.getvalue()


def _make_pdf() -> bytes:
    document = fitz.open()
    for text in ("ESB connection guide", "The timeout is 30 seconds"):
        page = document.new_page()
        page.insert_text((72, 72), text)
    content = document.tobytes()
    document.close()
    return content


def _make_xlsx() -> bytes:
    target = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Parameters"
    sheet.append(["parameter", "type", "default"])
    sheet.append(["timeout", "integer", 30])
    second = workbook.create_sheet("Errors")
    second.append(["code", "meaning"])
    second.append(["E001", "authentication failed"])
    workbook.save(target)
    workbook.close()
    return target.getvalue()


def _samples() -> tuple[Sample, ...]:
    markdown = (
        "# ESB Integration Guide\n\n"
        "Request credentials before connecting.\n\n"
        "## Parameters\n\n"
        "| parameter | default |\n| --- | --- |\n| timeout | 30 |\n"
    ).encode()
    return (
        Sample("guide.md", "text/markdown", markdown, SourceType.MARKDOWN),
        Sample(
            "guide.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _make_docx(),
            SourceType.DOCX,
        ),
        Sample("guide.pdf", "application/pdf", _make_pdf(), SourceType.PDF),
        Sample(
            "guide.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            _make_xlsx(),
            SourceType.XLSX,
        ),
    )


def _require_safe_environment(settings: Settings) -> None:
    database = make_url(settings.database_url).database or ""
    redis_database = urlparse(settings.redis_url).path
    bucket = settings.object_storage.bucket
    assert settings.environment == "integration"
    assert database == "knowagent_integration"
    assert redis_database == "/15"
    assert bucket == "knowagent-phase1-it"
    assert settings.object_storage.configured


def _s3_client(storage: ObjectStorageSettings) -> object:
    return boto3.client(
        "s3",
        endpoint_url=storage.endpoint_url,
        region_name=storage.region,
        aws_access_key_id=storage.access_key,
        aws_secret_access_key=storage.secret_key,
        verify=storage.verify_value,
    )


def _ensure_bucket(storage: ObjectStorageSettings) -> object:
    client = _s3_client(storage)
    buckets = {item["Name"] for item in client.list_buckets()["Buckets"]}
    if storage.bucket not in buckets:
        client.create_bucket(Bucket=storage.bucket)
    return client


def _bucket_keys(client: object, bucket: str) -> set[str]:
    keys: set[str] = set()
    paginator = client.get_paginator("list_objects_v2")
    for response in paginator.paginate(Bucket=bucket):
        keys.update(item["Key"] for item in response.get("Contents", []))
    return keys


def _delete_bucket_keys(client: object, bucket: str, keys: set[str]) -> None:
    objects = [{"Key": key} for key in sorted(keys)]
    if objects:
        client.delete_objects(Bucket=bucket, Delete={"Objects": objects})


def _restore_redis(
    redis: Redis,
    *,
    keys_before: set[str],
    queue_length_before: int,
) -> None:
    keys_after = set(redis.scan_iter())
    created_keys = keys_after - keys_before
    if created_keys:
        redis.delete(*created_keys)
    if "ingestion" in keys_before:
        added_messages = max(0, redis.llen("ingestion") - queue_length_before)
        if added_messages:
            redis.ltrim("ingestion", added_messages, -1)


def _require_status(response: object, expected: int) -> None:
    assert getattr(response, "status_code") == expected, getattr(response, "text")


def _login(client: TestClient, entry: str, username: str) -> dict[str, object]:
    response = client.post(
        f"/api/v1/auth/{entry}/sessions",
        json={"username": username, "password": PASSWORD},
    )
    _require_status(response, 200)
    return response.json()


def _process_job(
    *,
    app: object,
    store: S3ObjectStore,
    settings: Settings,
    job_id: UUID,
) -> ChunkManifest:
    result = IngestionProcessor(
        coordinator=app.state.ingestion_coordinator,
        object_store=store,
        parser_registry=ParserRegistry.default(
            ParserLimits.from_settings(settings.document_processing)
        ),
        chunker=StructureAwareChunker(ChunkingConfig.from_settings(settings.document_processing)),
        lease_seconds=settings.ingestion.lease_seconds,
        retry_base_seconds=1,
    ).process(job_id, worker_id=f"phase1-acceptance-{job_id}")
    assert result is not None
    assert result.job.status is IngestionStatus.SUCCEEDED
    assert result.version.status is DocumentVersionStatus.CHUNKED
    assert result.version.chunk_manifest_key is not None
    return ChunkManifest.model_validate_json(store.get(key=result.version.chunk_manifest_key))


def _publish_manifest(
    *,
    app: object,
    system_id: UUID,
    version_id: UUID,
    manifest: ChunkManifest,
    now: datetime,
) -> None:
    with app.state.session_factory.begin() as session:
        version = session.get(DocumentVersionRecord, version_id)
        assert version is not None
        version.status = DocumentVersionStatus.READY_DRAFT
    drafts = tuple(
        KnowledgeChunkDraft.model_validate(chunk.model_dump()) for chunk in manifest.chunks
    )
    assert drafts
    publication = KnowledgePublicationService(app.state.session_factory)
    publication.replace_draft_chunks(
        system_id=system_id,
        document_version_id=version_id,
        chunks=drafts,
        now=now,
    )
    publication.publish(
        system_id=system_id,
        document_version_id=version_id,
        now=now,
    )


def _assert_locator_contract(manifest: ChunkManifest, source_type: SourceType) -> None:
    assert manifest.source_type is source_type
    locators = [locator for chunk in manifest.chunks for locator in chunk.locators]
    assert locators
    if source_type is SourceType.MARKDOWN:
        assert any(locator.heading_path and locator.line_start for locator in locators)
    elif source_type is SourceType.DOCX:
        assert any(locator.heading_path and locator.paragraph_start for locator in locators)
        assert any(locator.table_index and locator.cell_range for locator in locators)
    elif source_type is SourceType.PDF:
        assert {locator.page_number for locator in locators} == {1, 2}
        assert all(locator.bounding_box is not None for locator in locators)
    elif source_type is SourceType.XLSX:
        assert {locator.sheet_name for locator in locators} == {"Parameters", "Errors"}
        assert all(locator.cell_range for locator in locators)


def _upload(  # pylint: disable=too-many-arguments
    *,
    client: TestClient,
    csrf_token: str,
    system_id: UUID,
    run_id: str,
    sample: Sample,
    suffix: str,
    document_id: UUID | None = None,
) -> object:
    data = {"document_name": f"Phase 1 {sample.source_type.value} guide"}
    if document_id is not None:
        data["document_id"] = str(document_id)
    response = client.post(
        f"/api/v1/systems/{system_id}/documents",
        headers={
            "X-CSRF-Token": csrf_token,
            "Idempotency-Key": f"phase1-{run_id}-{suffix}",
        },
        data=data,
        files={"file": (sample.filename, sample.content, sample.media_type)},
    )
    _require_status(response, 202)
    return response


def _verify_s3_contract(settings: Settings, raw_client: object, run_id: str) -> None:
    store = S3ObjectStore.from_settings(settings.object_storage)
    key = f"acceptance/{run_id}/multipart.bin"
    payload = b"phase1" * ((settings.object_storage.multipart_threshold // 6) + 1)
    store.put(
        key=key,
        content=BytesIO(payload),
        content_type="application/octet-stream",
        content_length=len(payload),
    )
    assert store.get(key=key) == payload
    assert "-" in raw_client.head_object(Bucket=settings.object_storage.bucket, Key=key)["ETag"]
    store.delete(key=key)
    with pytest.raises(ObjectStoreError) as missing:
        store.get(key=key)
    assert missing.value.retryable is False

    denied_store = S3ObjectStore.from_settings(
        replace(settings.object_storage, secret_key=f"wrong-{run_id}")
    )
    with pytest.raises(ObjectStoreError) as denied:
        denied_store.put(
            key=f"acceptance/{run_id}/denied.bin",
            content=BytesIO(b"denied"),
            content_type="application/octet-stream",
            content_length=6,
        )
    assert denied.value.retryable is False

    unavailable_store = S3ObjectStore.from_settings(
        replace(
            settings.object_storage,
            endpoint_url="http://127.0.0.1:1",
            connect_timeout_seconds=1,
            read_timeout_seconds=1,
            sdk_max_attempts=1,
        )
    )
    with pytest.raises(ObjectStoreError) as unavailable:
        unavailable_store.put(
            key=f"acceptance/{run_id}/unavailable.bin",
            content=BytesIO(b"retry"),
            content_type="application/octet-stream",
            content_length=5,
        )
    assert unavailable.value.retryable is True


def _seed_accounts(app: object, run_id: str) -> tuple[str, str]:
    password_hash = Argon2PasswordHasher().hash(PASSWORD)
    admin_name = f"it.admin.{run_id}"
    owner_name = f"it.owner.{run_id}"
    with app.state.session_factory.begin() as session:
        session.add_all(
            [
                AccountRecord(
                    username=admin_name,
                    display_name="Phase 1 Integration Admin",
                    password_hash=password_hash,
                    role=AccountRole.ADMIN,
                    source=AccountSource.ADMIN_CREATED,
                    status=AccountStatus.ACTIVE,
                    must_change_password=False,
                    session_version=1,
                ),
                AccountRecord(
                    username=owner_name,
                    display_name="Phase 1 Integration Owner",
                    password_hash=password_hash,
                    role=AccountRole.SYSTEM_OWNER,
                    source=AccountSource.ADMIN_CREATED,
                    status=AccountStatus.ACTIVE,
                    must_change_password=False,
                    session_version=1,
                ),
            ]
        )
    return admin_name, owner_name


def _cleanup_acceptance_records(app: object) -> None:
    with app.state.session_factory.begin() as session:
        system_ids = tuple(
            session.scalars(
                select(BusinessSystemRecord.id).where(
                    BusinessSystemRecord.name.in_(("Phase 1 System A", "Phase 1 System B"))
                )
            )
        )
        account_ids = tuple(
            session.scalars(
                select(AccountRecord.id).where(
                    AccountRecord.display_name.in_(
                        ("Phase 1 Integration Admin", "Phase 1 Integration Owner")
                    )
                )
            )
        )
        if system_ids:
            session.execute(delete(DocumentRecord).where(DocumentRecord.system_id.in_(system_ids)))
            session.execute(
                delete(AccountSystemRoleRecord).where(
                    AccountSystemRoleRecord.system_id.in_(system_ids)
                )
            )
            session.execute(
                delete(BusinessSystemRecord).where(BusinessSystemRecord.id.in_(system_ids))
            )
        if account_ids:
            session.execute(delete(AuditLogRecord).where(AuditLogRecord.actor_id.in_(account_ids)))
            session.execute(delete(AccountRecord).where(AccountRecord.id.in_(account_ids)))


def _acceptance_records_exist(app: object) -> bool:
    with app.state.session_factory() as session:
        system_id = session.scalar(
            select(BusinessSystemRecord.id)
            .where(BusinessSystemRecord.name.in_(("Phase 1 System A", "Phase 1 System B")))
            .limit(1)
        )
        account_id = session.scalar(
            select(AccountRecord.id)
            .where(
                AccountRecord.display_name.in_(
                    ("Phase 1 Integration Admin", "Phase 1 Integration Owner")
                )
            )
            .limit(1)
        )
    return system_id is not None or account_id is not None


def _create_systems_and_assign_owner(  # pylint: disable=too-many-arguments
    client: TestClient,
    *,
    admin_csrf: str,
    owner_name: str,
    run_id: str,
) -> tuple[UUID, UUID]:
    responses = [
        client.post(
            "/api/v1/admin/systems",
            headers={"X-CSRF-Token": admin_csrf},
            json={"code": f"A{run_id}", "name": "Phase 1 System A"},
        ),
        client.post(
            "/api/v1/admin/systems",
            headers={"X-CSRF-Token": admin_csrf},
            json={"code": f"B{run_id}", "name": "Phase 1 System B"},
        ),
    ]
    for response in responses:
        _require_status(response, 201)
    system_a, system_b = (UUID(response.json()["id"]) for response in responses)
    accounts = client.get(
        "/api/v1/admin/accounts",
        params={"role": "SYSTEM_OWNER", "page_size": 100},
    )
    _require_status(accounts, 200)
    owner_id = next(
        UUID(item["id"]) for item in accounts.json()["items"] if item["username"] == owner_name
    )
    assigned = client.put(
        f"/api/v1/admin/systems/{system_a}/owners",
        headers={"X-CSRF-Token": admin_csrf},
        json={"account_ids": [str(owner_id)], "replace_existing": True},
    )
    _require_status(assigned, 200)
    return system_a, system_b


def _verify_publication_isolation(  # pylint: disable=too-many-arguments
    app: object,
    *,
    system_a: UUID,
    system_b: UUID,
    markdown_document_id: UUID,
    markdown_v1: UUID,
    markdown_v2: UUID,
) -> None:
    with app.state.session_factory() as session:
        knowledge = SqlAlchemyKnowledgeRepository(session)
        document = knowledge.get_document(system_id=system_a, document_id=markdown_document_id)
        first = knowledge.get_version(system_id=system_a, document_version_id=markdown_v1)
        second = knowledge.get_version(system_id=system_a, document_version_id=markdown_v2)
        published_a = knowledge.list_published_chunks(system_id=system_a, limit=100)
        published_b = knowledge.list_published_chunks(system_id=system_b, limit=100)
        assert document is not None
        assert document.current_published_version_id == markdown_v2
        assert first is not None and first.publish_status is PublicationStatus.RETIRED
        assert second is not None and second.publish_status is PublicationStatus.PUBLISHED
        assert any("SYSTEM-A-V2-ONLY" in chunk.text for chunk in published_a)
        assert all("SYSTEM-B-ONLY" not in chunk.text for chunk in published_a)
        assert any("SYSTEM-B-ONLY" in chunk.text for chunk in published_b)
        assert all("SYSTEM-A-V2-ONLY" not in chunk.text for chunk in published_b)
        assert knowledge.get_document(system_id=system_b, document_id=markdown_document_id) is None


def _verify_recovery(  # pylint: disable=too-many-arguments
    app: object,
    *,
    settings: Settings,
    owner_client: TestClient,
    owner_csrf: str,
    system_id: UUID,
    run_id: str,
) -> tuple[UUID, int]:
    sample = Sample("recovery.md", "text/markdown", b"# Recovery\n", SourceType.MARKDOWN)
    response = _upload(
        client=owner_client,
        csrf_token=owner_csrf,
        system_id=system_id,
        run_id=run_id,
        sample=sample,
        suffix="recovery",
    )
    job_id = UUID(response.json()["job_id"])
    claim_time = datetime.now(UTC)
    claimed = app.state.ingestion_coordinator.claim(
        job_id,
        owner="terminated-worker",
        now=claim_time,
        lease_seconds=1,
    )
    assert claimed is not None and claimed.job.status is IngestionStatus.RUNNING
    recovered = IngestionRecoveryService(
        coordinator=app.state.ingestion_coordinator,
        dispatcher=CeleryIngestionDispatcher(build_celery_app(settings)),
        dispatch_stale_seconds=1,
        batch_size=10,
        clock=lambda: claim_time + timedelta(seconds=2),
    ).run()
    assert recovered == 1
    with app.state.session_factory() as session:
        bundle = SqlAlchemyIngestionRepository(session).get_by_job_id(job_id)
        assert bundle is not None
        assert bundle.job.status is IngestionStatus.QUEUED
        assert bundle.job.celery_task_id is not None
    return job_id, recovered


def test_phase1_live_acceptance() -> None:  # pylint: disable=too-many-locals,too-many-statements
    settings = Settings.from_environment()
    _require_safe_environment(settings)
    run_id = uuid4().hex[:8]
    redis = Redis.from_url(settings.redis_url, decode_responses=True)
    redis_keys_before = set(redis.scan_iter())
    queue_length_before = redis.llen("ingestion")
    raw_s3 = _ensure_bucket(settings.object_storage)
    s3_keys_before = _bucket_keys(raw_s3, settings.object_storage.bucket)
    store = S3ObjectStore.from_settings(settings.object_storage)
    app = create_app(settings)
    app.state.object_store = store
    _cleanup_acceptance_records(app)
    admin_name, owner_name = _seed_accounts(app, run_id)
    accepted = False

    try:
        _verify_s3_contract(settings, raw_s3, run_id)
        with (
            TestClient(app, base_url="http://testserver") as admin_client,
            TestClient(app, base_url="http://testserver") as owner_client,
        ):
            admin_session = _login(admin_client, "admin", admin_name)
            system_a, system_b = _create_systems_and_assign_owner(
                admin_client,
                admin_csrf=str(admin_session["csrf_token"]),
                owner_name=owner_name,
                run_id=run_id,
            )
            owner_session = _login(owner_client, "user", owner_name)
            owner_csrf = str(owner_session["csrf_token"])

            manifests: dict[SourceType, ChunkManifest] = {}
            versions: dict[SourceType, UUID] = {}
            documents: dict[SourceType, UUID] = {}
            markdown_response: object | None = None
            for index, sample in enumerate(_samples()):
                response = _upload(
                    client=owner_client,
                    csrf_token=owner_csrf,
                    system_id=system_a,
                    run_id=run_id,
                    sample=sample,
                    suffix=f"{index}-v1",
                )
                if sample.source_type is SourceType.MARKDOWN:
                    replay = _upload(
                        client=owner_client,
                        csrf_token=owner_csrf,
                        system_id=system_a,
                        run_id=run_id,
                        sample=sample,
                        suffix=f"{index}-v1",
                    )
                    assert replay.json()["job_id"] == response.json()["job_id"]
                    markdown_response = response
                document_id = UUID(response.json()["document_id"])
                version_id = UUID(response.json()["document_version_id"])
                manifest = _process_job(
                    app=app,
                    store=store,
                    settings=settings,
                    job_id=UUID(response.json()["job_id"]),
                )
                _assert_locator_contract(manifest, sample.source_type)
                _publish_manifest(
                    app=app,
                    system_id=system_a,
                    version_id=version_id,
                    manifest=manifest,
                    now=datetime.now(UTC),
                )
                manifests[sample.source_type] = manifest
                versions[sample.source_type] = version_id
                documents[sample.source_type] = document_id

            assert markdown_response is not None
            forbidden = owner_client.post(
                f"/api/v1/systems/{system_b}/documents",
                headers={
                    "X-CSRF-Token": owner_csrf,
                    "Idempotency-Key": f"phase1-{run_id}-forbidden",
                },
                data={"document_name": "Forbidden"},
                files={"file": ("forbidden.md", b"# Forbidden\n", "text/markdown")},
            )
            _require_status(forbidden, 403)
            assert forbidden.json()["code"] == "SYSTEM_ACCESS_DENIED"

            markdown_v2_sample = Sample(
                "guide-v2.md",
                "text/markdown",
                b"# ESB Guide\n\nSYSTEM-A-V2-ONLY\n",
                SourceType.MARKDOWN,
            )
            markdown_v2_response = _upload(
                client=owner_client,
                csrf_token=owner_csrf,
                system_id=system_a,
                run_id=run_id,
                sample=markdown_v2_sample,
                suffix="markdown-v2",
                document_id=documents[SourceType.MARKDOWN],
            )
            assert markdown_v2_response.json()["version_no"] == 2
            markdown_v2 = UUID(markdown_v2_response.json()["document_version_id"])
            markdown_v2_manifest = _process_job(
                app=app,
                store=store,
                settings=settings,
                job_id=UUID(markdown_v2_response.json()["job_id"]),
            )
            _publish_manifest(
                app=app,
                system_id=system_a,
                version_id=markdown_v2,
                manifest=markdown_v2_manifest,
                now=datetime.now(UTC) + timedelta(seconds=1),
            )

            system_b_sample = Sample(
                "system-b.md",
                "text/markdown",
                b"# ESB Guide\n\nSYSTEM-B-ONLY\n",
                SourceType.MARKDOWN,
            )
            system_b_response = _upload(
                client=admin_client,
                csrf_token=str(admin_session["csrf_token"]),
                system_id=system_b,
                run_id=run_id,
                sample=system_b_sample,
                suffix="system-b",
            )
            system_b_manifest = _process_job(
                app=app,
                store=store,
                settings=settings,
                job_id=UUID(system_b_response.json()["job_id"]),
            )
            _publish_manifest(
                app=app,
                system_id=system_b,
                version_id=UUID(system_b_response.json()["document_version_id"]),
                manifest=system_b_manifest,
                now=datetime.now(UTC),
            )
            _verify_publication_isolation(
                app,
                system_a=system_a,
                system_b=system_b,
                markdown_document_id=documents[SourceType.MARKDOWN],
                markdown_v1=versions[SourceType.MARKDOWN],
                markdown_v2=markdown_v2,
            )
            recovery_job, recovered = _verify_recovery(
                app,
                settings=settings,
                owner_client=owner_client,
                owner_csrf=owner_csrf,
                system_id=system_a,
                run_id=run_id,
            )

        session_keys = list(redis.scan_iter(match=f"{settings.redis_prefix}:*"))
        assert session_keys
        assert redis.llen("ingestion") >= 1
        print(
            json.dumps(
                {
                    "formats": sorted(source.value for source in manifests),
                    "idempotency_replayed": True,
                    "multipart_verified": True,
                    "permission_denied_verified": True,
                    "recovered_job_id": str(recovery_job),
                    "recovered_jobs_dispatched": recovered,
                    "redis_session_keys": len(session_keys),
                    "systems": 2,
                    "version_switch": "v1-retired-v2-published",
                    "zero_cross_system_leakage": True,
                },
                sort_keys=True,
            )
        )
        accepted = True
    finally:
        app.state.redis_client.close()
        if accepted:
            _cleanup_acceptance_records(app)
            assert not _acceptance_records_exist(app)
        app.state.engine.dispose()
        _restore_redis(
            redis,
            keys_before=redis_keys_before,
            queue_length_before=queue_length_before,
        )
        assert set(redis.scan_iter()) == redis_keys_before
        assert redis.llen("ingestion") == queue_length_before
        redis.close()
        _delete_bucket_keys(
            raw_s3,
            settings.object_storage.bucket,
            _bucket_keys(raw_s3, settings.object_storage.bucket) - s3_keys_before,
        )
        assert _bucket_keys(raw_s3, settings.object_storage.bucket) == s3_keys_before
