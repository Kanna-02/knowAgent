from __future__ import annotations

from io import BytesIO
from uuid import UUID, uuid4
from zipfile import ZIP_DEFLATED, ZipFile

import fitz
import pytest
from docx import Document
from openpyxl import Workbook

from knowagent.documents.domain.models import BlockType, ParsedDocument, SourceType
from knowagent.documents.errors import DocumentParseError, ParseErrorCode
from knowagent.documents.infrastructure.parsers import (
    DocxDocumentParser,
    MarkdownDocumentParser,
    ParserLimits,
    ParserRegistry,
    PdfDocumentParser,
    XlsxDocumentParser,
)
from knowagent.documents.ports import DocumentParser


def parse(
    parser: DocumentParser,
    content: bytes,
    *,
    document_id: UUID | None = None,
    document_version_id: UUID | None = None,
) -> ParsedDocument:
    return parser.parse(
        content=content,
        document_id=document_id or uuid4(),
        document_version_id=document_version_id or uuid4(),
    )


def make_docx() -> bytes:
    target = BytesIO()
    document = Document()
    document.add_heading("接入指南", level=1)
    document.add_paragraph("先申请客户端凭据。")
    document.add_heading("超时", level=2)
    document.add_paragraph("默认超时为 30 秒。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "参数"
    table.cell(0, 1).text = "含义"
    table.cell(1, 0).text = "timeout"
    table.cell(1, 1).text = "请求超时"
    document.save(target)
    return target.getvalue()


def make_pdf(*pages: str) -> bytes:
    document = fitz.open()
    for text in pages:
        page = document.new_page()
        if text:
            page.insert_text((72, 72), text)
    content = document.tobytes()
    document.close()
    return content


def make_encrypted_pdf() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "protected")
    content = document.tobytes(
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner-secret",
        user_pw="user-secret",
    )
    document.close()
    return content


def make_xlsx() -> bytes:
    target = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "参数"
    sheet.append(["参数", "类型", "默认值"])
    sheet.append(["timeout", "integer", 30])
    sheet.append(["endpoint", "string", "https://example.invalid"])
    second = workbook.create_sheet("错误码")
    second.append(["代码", "说明"])
    second.append(["E001", "鉴权失败"])
    workbook.save(target)
    workbook.close()
    return target.getvalue()


def corrupt_archive_member(content: bytes, member_name: str) -> bytes:
    target = BytesIO()
    with ZipFile(BytesIO(content)) as source, ZipFile(target, "w", ZIP_DEFLATED) as output:
        for member in source.infolist():
            value = b"<broken" if member.filename == member_name else source.read(member.filename)
            output.writestr(member, value)
    return target.getvalue()


def test_docx_parser_preserves_heading_paragraph_and_table_locations() -> None:
    parsed = parse(DocxDocumentParser(), make_docx())

    assert parsed.source_type is SourceType.DOCX
    assert [block.block_index for block in parsed.blocks] == list(range(len(parsed.blocks)))
    timeout = next(block for block in parsed.blocks if "默认超时" in block.text)
    assert timeout.locator.heading_path == ("接入指南", "超时")
    assert timeout.locator.paragraph_start == timeout.locator.paragraph_end
    table_rows = [block for block in parsed.blocks if block.block_type is BlockType.TABLE_ROW]
    assert [block.text for block in table_rows] == ["参数 | 含义", "timeout | 请求超时"]
    assert table_rows[1].locator.table_index == 1
    assert table_rows[1].locator.cell_range == "A2:B2"


def test_markdown_parser_preserves_ast_structure_and_source_lines() -> None:
    content = (
        "# 接入指南\n\n"
        "申请凭据后调用接口。\n\n"
        "## 参数\n\n"
        "- `timeout` 必填\n"
        "- `retries` 可选\n\n"
        '```json\n{"timeout": 30}\n```\n\n'
        "| 参数 | 默认值 |\n| --- | --- |\n| timeout | 30 |\n"
    ).encode()

    parsed = parse(MarkdownDocumentParser(), content)

    paragraph = next(block for block in parsed.blocks if "申请凭据" in block.text)
    assert paragraph.locator.heading_path == ("接入指南",)
    assert (paragraph.locator.line_start, paragraph.locator.line_end) == (3, 3)
    assert any(block.block_type is BlockType.LIST_ITEM for block in parsed.blocks)
    assert any(block.block_type is BlockType.CODE for block in parsed.blocks)
    table_rows = [block for block in parsed.blocks if block.block_type is BlockType.TABLE_ROW]
    assert [block.text for block in table_rows] == ["参数 | 默认值", "timeout | 30"]


def test_markdown_parser_maps_each_table_row_to_its_exact_source_line() -> None:
    content = (
        "| 参数 | 默认值 |\n"
        "| --- | --- |\n"
        "| timeout | 30 |\n"
        "| retries | 3 |\n"
        "| enabled | true |\n"
    ).encode()

    parsed = parse(MarkdownDocumentParser(), content)
    table_rows = [block for block in parsed.blocks if block.block_type is BlockType.TABLE_ROW]

    assert [block.locator.line_start for block in table_rows] == [1, 3, 4, 5]
    assert [block.locator.line_end for block in table_rows] == [1, 3, 4, 5]


def test_pdf_parser_preserves_page_and_bounding_box() -> None:
    parsed = parse(PdfDocumentParser(), make_pdf("ESB connection guide", "Timeout is 30 seconds"))

    assert {block.locator.page_number for block in parsed.blocks} == {1, 2}
    assert all(block.locator.bounding_box is not None for block in parsed.blocks)
    assert "Timeout" in parsed.blocks[-1].text


def test_pdf_parser_marks_empty_or_scanned_document_for_ocr() -> None:
    with pytest.raises(DocumentParseError) as caught:
        parse(PdfDocumentParser(), make_pdf(""))

    assert caught.value.code is ParseErrorCode.OCR_REQUIRED


def test_pdf_parser_rejects_corrupt_and_password_protected_files() -> None:
    for content, expected_code in [
        (b"not-a-pdf", ParseErrorCode.INVALID_FILE),
        (make_encrypted_pdf(), ParseErrorCode.PASSWORD_PROTECTED),
    ]:
        with pytest.raises(DocumentParseError) as caught:
            parse(PdfDocumentParser(), content)
        assert caught.value.code is expected_code


def test_xlsx_parser_preserves_sheet_and_exact_row_ranges() -> None:
    parsed = parse(XlsxDocumentParser(), make_xlsx())

    assert parsed.source_type is SourceType.XLSX
    assert [block.locator.sheet_name for block in parsed.blocks] == [
        "参数",
        "参数",
        "参数",
        "错误码",
        "错误码",
    ]
    assert parsed.blocks[1].locator.cell_range == "A2:C2"
    assert parsed.blocks[1].text == "timeout | integer | 30"
    assert parsed.blocks[3].table_header is True


def test_parser_registry_selects_by_extension_and_media_type() -> None:
    registry = ParserRegistry.default()

    assert (
        registry.resolve(filename="guide.PDF", media_type="application/pdf").source_type
        is SourceType.PDF
    )
    assert (
        registry.resolve(filename="guide.md", media_type="text/markdown").source_type
        is SourceType.MARKDOWN
    )
    assert (
        registry.resolve(filename="guide.md", media_type="text/plain; charset=utf-8").source_type
        is SourceType.MARKDOWN
    )
    with pytest.raises(DocumentParseError) as caught:
        registry.resolve(filename="legacy.doc", media_type="application/msword")
    assert caught.value.code is ParseErrorCode.UNSUPPORTED_FORMAT


def test_parsers_report_corrupt_files_and_resource_limits() -> None:
    with pytest.raises(DocumentParseError) as corrupt:
        parse(DocxDocumentParser(), b"not-a-zip")
    assert corrupt.value.code is ParseErrorCode.INVALID_FILE

    limits = ParserLimits(max_pdf_pages=1)
    with pytest.raises(DocumentParseError) as limited:
        parse(PdfDocumentParser(limits), make_pdf("first page", "second page"))
    assert limited.value.code is ParseErrorCode.RESOURCE_LIMIT_EXCEEDED


def test_docx_parser_enforces_block_limit_and_detects_encrypted_container() -> None:
    with pytest.raises(DocumentParseError) as limited:
        parse(DocxDocumentParser(ParserLimits(max_docx_blocks=1)), make_docx())
    assert limited.value.code is ParseErrorCode.RESOURCE_LIMIT_EXCEEDED

    encrypted_signature = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"encrypted"
    with pytest.raises(DocumentParseError) as encrypted:
        parse(DocxDocumentParser(), encrypted_signature)
    assert encrypted.value.code is ParseErrorCode.PASSWORD_PROTECTED

    with pytest.raises(DocumentParseError) as too_many_members:
        parse(DocxDocumentParser(ParserLimits(max_archive_members=1)), make_docx())
    assert too_many_members.value.code is ParseErrorCode.RESOURCE_LIMIT_EXCEEDED


def test_markdown_parser_rejects_invalid_encoding_and_file_size_limit() -> None:
    for content, limits, expected_code in [
        (b"\xff\xfe\xfd", ParserLimits(), ParseErrorCode.DECODING_ERROR),
        (b"# too large", ParserLimits(max_file_bytes=2), ParseErrorCode.RESOURCE_LIMIT_EXCEEDED),
    ]:
        with pytest.raises(DocumentParseError) as caught:
            parse(MarkdownDocumentParser(limits), content)
        assert caught.value.code is expected_code

    with pytest.raises(DocumentParseError) as too_many_blocks:
        parse(
            MarkdownDocumentParser(ParserLimits(max_markdown_blocks=1)),
            b"# Heading\n\nSecond block\n",
        )
    assert too_many_blocks.value.code is ParseErrorCode.RESOURCE_LIMIT_EXCEEDED


def test_xlsx_parser_rejects_corrupt_file_and_sheet_limit() -> None:
    with pytest.raises(DocumentParseError) as corrupt:
        parse(XlsxDocumentParser(), b"not-an-xlsx")
    assert corrupt.value.code is ParseErrorCode.INVALID_FILE

    with pytest.raises(DocumentParseError) as limited:
        parse(XlsxDocumentParser(ParserLimits(max_xlsx_sheets=1)), make_xlsx())
    assert limited.value.code is ParseErrorCode.RESOURCE_LIMIT_EXCEEDED


@pytest.mark.parametrize(
    ("parser", "content"),
    [
        (
            DocxDocumentParser(),
            corrupt_archive_member(make_docx(), "word/document.xml"),
        ),
        (
            XlsxDocumentParser(),
            corrupt_archive_member(make_xlsx(), "xl/workbook.xml"),
        ),
    ],
    ids=["docx", "xlsx"],
)
def test_office_parsers_map_malformed_internal_xml_to_invalid_file(
    parser: DocumentParser, content: bytes
) -> None:
    with pytest.raises(DocumentParseError) as caught:
        parse(parser, content)

    assert caught.value.code is ParseErrorCode.INVALID_FILE


@pytest.mark.parametrize(
    ("parser", "content"),
    [
        (DocxDocumentParser(), make_docx()),
        (MarkdownDocumentParser(), b"# Guide\n\nStable content\n"),
        (PdfDocumentParser(), make_pdf("Stable content")),
        (XlsxDocumentParser(), make_xlsx()),
    ],
    ids=["docx", "markdown", "pdf", "xlsx"],
)
def test_parser_rerun_with_same_ids_is_idempotent(parser: DocumentParser, content: bytes) -> None:
    document_id = uuid4()
    version_id = uuid4()

    first = parse(
        parser,
        content,
        document_id=document_id,
        document_version_id=version_id,
    )
    second = parse(
        parser,
        content,
        document_id=document_id,
        document_version_id=version_id,
    )

    assert first == second
